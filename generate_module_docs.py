from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4');    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 112, 192))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def heading3(text, color=(68,68,68)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=color)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)

def body(text, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(20, 20, 20)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F4F4F8')
        pPr.append(shd)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

def info_box(lines, bg='EBF3FB'):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=10)
        run.font.color.rgb = RGBColor(20, 20, 80)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), bg)
        pPr.append(shd)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

def warn_box(lines):
    info_box(lines, bg='FFF3CD')

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255,255,255))
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'2E74B5')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        fill = 'FFFFFF' if ri%2==0 else 'EBF3FB'
        for ci, val in enumerate(row):
            cell = table.cell(ri+1, ci)
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=9.5)
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AJ SHRINK ERP")
set_font(run, size=28, bold=True, color=(31,73,125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Module Documentation")
set_font(run, size=18, color=(0,112,192))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Product Catalog  |  Gravure Estimation  |  Cost Estimation")
set_font(run, size=13, italic=True, color=(80,80,80))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0  |  April 2026  |  AJ Shrink Flexible Packaging ERP")
set_font(run, size=10, color=(130,130,130))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 1 — OVERVIEW (Teen Modules)
# ══════════════════════════════════════════════════════════════

heading1("1. Teen Modules Ka Overview")

body("Is documentation mein teen connected modules cover hain jo gravure printing ke liye "
     "end-to-end workflow banate hain:")

add_table(
    ["Module", "Route (URL)", "File", "Kaam"],
    [
        ["Product Catalog",    "/gravure/product-catalog", "app/gravure/product-catalog/page.tsx", "Product ka technical specification + cylinder planning + material prep"],
        ["Gravure Estimation", "/gravure/estimation",      "app/gravure/estimation/page.tsx",       "Costing + production plan + pricing calculation"],
        ["Cost Estimation",    "/cost-estimation",         "app/cost-estimation/page.tsx",          "Extrusion film recipe-based raw material costing"],
    ],
    col_widths=[4, 4.5, 6, 7]
)

heading2("1.1  Workflow — Kaise Connected Hain?")
code_block([
    "Enquiry  →  Gravure Order  →  Product Catalog  →  Estimation  →  Order Confirmation",
    "",
    "  1. Customer enquiry aata hai  (app/gravure/enquiry)",
    "  2. Order book hota hai        (app/gravure/orders)",
    "  3. Product Catalog entry banati hai — cylinder plan, material, color shades",
    "  4. Estimation page cost calculate karta hai — per meter rate nikalta hai",
    "  5. Cost Estimation — film recipe se raw material cost (extrusion ke liye)",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 2 — PRODUCT CATALOG
# ══════════════════════════════════════════════════════════════

heading1("2. Product Catalog — Poori Detail")

body("File: app/gravure/product-catalog/page.tsx")
body("Route: /gravure/product-catalog")

heading2("2.1  Yeh Page Kya Karta Hai?")
body("Product Catalog ek registered product database hai. Jab bhi koi naya product gravure printing "
     "ke liye approve hota hai, uska poora specification — dimensions, ply structure, cylinder plan, "
     "color shades — yahan save hota hai. Ek baar save ho gaya toh future orders mein seedha "
     "yahi specification reuse karte hain.")

heading2("2.2  Page Ke Do Sections (Tabs)")
add_table(
    ["Tab", "Naam", "Kya Dikhta Hai"],
    [
        ["Tab 1", "Pending",   "Gravure orders jinke liye abhi catalog entry nahi bani — yahan se naya catalog create karte hain"],
        ["Tab 2", "Processed", "Jo catalog entries ban chuki hain — view, edit (replan), delete kar sakte hain"],
    ],
    col_widths=[2, 3, 13]
)

heading2("2.3  Naya Catalog Entry Kaise Banate Hain?")

heading3("Step 1 — Pending tab se order select karo")
bullet("Pending orders ki list dikhi hogi")
bullet("Kisi order ke saamne 'Create Catalog' button click karo")
bullet("Ya fir 'Create Direct' button se bhi bana sakte hain bina order ke")

heading3("Step 2 — 3-Tab Modal khulta hai")
add_table(
    ["Modal Tab", "Naam", "Kya Fill Karna Hai"],
    [
        ["Tab 1", "Info",     "Product name, customer, category, job dimensions (W×H), colors, machine, standard qty"],
        ["Tab 2", "Planning", "Production plan select karo — cylinder circumference, UPS (across × repeat), raw material meters"],
        ["Tab 3", "Material", "Ply structure (film layers), consumables (ink, adhesive), color shades, cylinder allocation"],
    ],
    col_widths=[2.5, 3, 12.5]
)

heading2("2.4  Tab 1 — Info Section Detail")
add_table(
    ["Field", "Matlab", "Type"],
    [
        ["Product Name",    "Product ka naam jo customer ne diya",                  "Text input"],
        ["Customer",        "Customer dropdown",                                     "Select"],
        ["Category",        "Product category (Categories Master se aata hai)",      "Select"],
        ["Content Type",    "Label / Sleeve / Pouch / Multi-Pack Shrink",            "Auto-detect from category"],
        ["Job Width (mm)",  "Pouch/label ki actual width in millimeters",            "Number"],
        ["Job Height (mm)", "Pouch/label ki actual height",                          "Number"],
        ["Front Colors",    "Front side pe kitne colors hain",                       "Number"],
        ["Back Colors",     "Back side pe kitne colors hain",                        "Number"],
        ["No. of Colors",   "Front + Back (auto-calculate hota hai)",                "Auto"],
        ["Machine",         "Printing machine select karo",                          "Select from Machine Master"],
        ["Standard Qty",    "Ek run mein kitna quantity standard hai (meters)",      "Number"],
        ["Trimming Size",   "Trimming ke baad actual size adjustment (mm)",          "Number"],
        ["Width Shrinkage", "Sleeve ke liye shrinkage % lagta hai",                  "Number (%)"],
        ["Rate/Meter",      "Per meter selling rate",                                "Number"],
        ["Remark",          "Koi additional note",                                   "Textarea"],
    ],
    col_widths=[4.5, 9, 4.5]
)

heading2("2.5  Tab 2 — Planning Section (Cylinder Planning)")

body("Yeh section sabse important hai. Yahan production plan select karte hain jisme "
     "cylinder circumference aur UPS (Units Per Sheet) decide hota hai.")

heading3("Structure Types aur Planning Logic:")
add_table(
    ["Structure Type", "Kab Use Hota", "Film Width Calculation"],
    [
        ["Label",           "Normal labels, roll form",          "actualWidth use hota hai (trimming-adjusted)"],
        ["Sleeve",          "Shrink sleeves",                    "layflat = jobWidth; cylinder circ = layflat × 2"],
        ["Pouch",           "Stand-up pouch, zipper pouch",      "jobWidth as-is use hota hai"],
        ["MultiPackShrink", "LDPE multi-pack shrink film",       "Special calculation — pack width + margins"],
    ],
    col_widths=[4, 5, 9]
)

heading3("UPS (Units Per Sheet) Kya Hota Hai?")
body("UPS matlab hai ek cylinder revolution mein kitne products print hote hain.")
code_block([
    "UPS = Across UPS × Repeat UPS",
    "",
    "Across UPS  = Machine film width ÷ product width (ek row mein kitne across)",
    "Repeat UPS  = Cylinder circumference ÷ product height (ek revolution mein kitne repeat)",
    "",
    "Example:",
    "  Machine width    = 1300 mm",
    "  Product width    = 200 mm  →  Across UPS = 1300 ÷ 200 = 6",
    "  Cylinder circ    = 600 mm",
    "  Product height   = 200 mm  →  Repeat UPS = 600 ÷ 200 = 3",
    "  Total UPS        = 6 × 3 = 18  (ek revolution mein 18 labels)",
])

heading3("Production Plan Table:")
body("System automatically saare possible plans generate karta hai based on available "
     "cylinders aur machine constraints. Table mein:")
bullet("Cylinder circumference options (machine ke min-max range mein)")
bullet("Sleeve tools ka printWidth × 2 = circumference")
bullet("Har plan ke liye: Across UPS, Repeat UPS, Total UPS, Required Meters")
bullet("Best plan select karo — woh 'Applied' ho jaata hai")

heading2("2.6  Tab 3 — Material Section")

heading3("A) Ply Configuration (Film Layers)")
body("Product ki film structure define karte hain. Har ply ek layer hai:")
add_table(
    ["Field", "Matlab"],
    [
        ["Ply Type",       "Film / Ink Layer / Adhesive Layer / Lamination / Coating"],
        ["Film Sub-Group", "Film ki type: BOPP, PET, LLDPE, etc."],
        ["GSM",            "Grams per square meter — film ka weight"],
        ["Thickness (µ)",  "Film ki thickness microns mein (GSM se auto-calculate)"],
        ["Film Rate",      "Per kg rate (Estimation Rate Master se aata hai)"],
        ["Consumable Items", "Is ply ke saath ink, adhesive, hardner, solvent add karo"],
    ],
    col_widths=[4, 14]
)

heading3("B) Consumable Items (Ink, Adhesive, etc.)")
add_table(
    ["Field", "Matlab"],
    [
        ["Item Group",    "Ink / Adhesive / Hardner / Solvent"],
        ["Item Name",     "Specific item select karo Item Master se"],
        ["GSM",           "Coverage weight per m²"],
        ["Coverage %",    "Agar ink poore area pe nahi, toh coverage % lagao"],
        ["Rate",          "Per kg rate"],
    ],
    col_widths=[4, 14]
)

heading3("C) Shade Card / Color Shades Tab")
add_table(
    ["Field", "Matlab"],
    [
        ["Color No.",      "Color ka sequence number"],
        ["Color Name",     "Color ka naam (e.g., Cyan, Magenta, Yellow, Black)"],
        ["Ink Type",       "Spot / Process / Special"],
        ["Pantone Ref",    "Pantone color reference number"],
        ["LAB Values",     "L*, a*, b* — color measurement values"],
        ["Delta E",        "Color difference — actual vs standard"],
        ["Shade Card Ref", "Shade card ka reference number"],
        ["Status",         "Pending / Standard Received / Approved / Rejected"],
    ],
    col_widths=[4, 14]
)

heading3("D) Cylinder Allocation Tab")
add_table(
    ["Field", "Matlab"],
    [
        ["Cylinder No",     "Cylinder ka unique number"],
        ["Circumference",   "Cylinder ki circumference (mm)"],
        ["Print Width",     "Cylinder ki print width (mm)"],
        ["Cylinder Type",   "New / Existing / Rechromed / Repeat"],
        ["Status",          "Pending / Available / In Use / Under Chrome / Ordered"],
        ["Repeat UPS",      "Ek cylinder revolution mein repeat count"],
    ],
    col_widths=[4, 14]
)

heading2("2.7  Live Cost Preview (Material + Process)")
body("Catalog entry mein ek live cost preview bhi hota hai jo real-time calculate hota hai:")
code_block([
    "Material Cost   = Film Cost + Consumable Cost",
    "  Film Cost     = (GSM ÷ 1000) × Area m² × Rate per Kg",
    "  Consumable    = (Effective GSM ÷ 1000) × Area m² × Rate  (coverage % apply hota hai)",
    "",
    "Process Cost    = Process Rate × Qty (SQM/Meter/Color/Job)",
    "Cylinder Cost   = No. of Colors × Cost per Color",
    "Overhead        = Subtotal × Overhead %",
    "Profit          = (Subtotal + Overhead) × Profit %",
    "Per Meter Rate  = Total ÷ Standard Qty",
])

heading2("2.8  Attachments")
body("Catalog entry ke saath files attach kar sakte hain — design files, shade cards, artwork, etc.:")
bullet("Pehli attachment automatic 'Master File' label milti hai")
bullet("Har attachment ka label edit kar sakte hain")
bullet("PDF, image, DXF, AI files support hain")
bullet("Attachment preview available hai inline")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 3 — GRAVURE ESTIMATION
# ══════════════════════════════════════════════════════════════

heading1("3. Gravure Estimation — Poori Detail")

body("File: app/gravure/estimation/page.tsx")
body("Route: /gravure/estimation")

heading2("3.1  Yeh Page Kya Karta Hai?")
body("Gravure Estimation page ek complete costing tool hai. Yahan har job ke liye "
     "per-meter rate calculate karte hain jo customer ko quote kia jaata hai. "
     "Multiple quantities ka comparison bhi kar sakte hain.")

heading2("3.2  Estimation Status Lifecycle")
add_table(
    ["Status", "Matlab", "Color"],
    [
        ["Draft",    "Abhi bana hai, incomplete",                    "Grey"],
        ["Approved", "Internally approve ho gaya",                    "Green"],
        ["Sent",     "Customer ko bhej diya",                         "Blue"],
        ["Accepted", "Customer ne accept kar liya — order banega",    "Emerald"],
        ["Rejected", "Customer ne reject kar diya",                   "Red"],
    ],
    col_widths=[3, 9, 3]
)

heading2("3.3  Teen Tabs — Estimation Modal")
add_table(
    ["Tab No.", "Naam", "Kya Karna Hai"],
    [
        ["Tab 1", "Basic Info",            "Job details, dimensions, ply config, machine selection"],
        ["Tab 2", "View Plan (Production)", "Production plan select karo — UPS, cylinder, required meters"],
        ["Tab 3", "Cost Estimation",        "Material/process costing, multiple qty comparison, per-meter rate"],
    ],
    col_widths=[2.5, 5, 10.5]
)

heading2("3.4  Tab 1 — Basic Info Detail")

heading3("Sub-section: Identification & Category")
add_table(
    ["Field", "Matlab", "Source"],
    [
        ["Estimation No",  "Auto-generate hota hai (e.g., GE-2425-0001)",      "generateCode() function"],
        ["Date",           "Estimation ki date",                                 "Auto current date"],
        ["Category",       "Product category",                                   "Categories Master (live)"],
        ["Content Type",   "Label/Sleeve/Pouch — category se auto-set hota hai","Category Master"],
        ["Enquiry No",     "Linked enquiry (optional)",                          "Enquiry module se"],
        ["Customer",       "Customer naam",                                       "Customers list"],
        ["Job Name",       "Product/job ka naam",                                 "Manual"],
        ["Sales Person",   "Sales representative",                                "Manual"],
        ["Sales Type",     "Local / Export",                                      "Dropdown"],
    ],
    col_widths=[4, 8, 6]
)

heading3("Sub-section: Planning Specification (Dimensions)")
add_table(
    ["Field", "Matlab", "Unit"],
    [
        ["Job Width",        "Product ki width",                    "mm"],
        ["Job Height",       "Product ki height / repeat length",   "mm"],
        ["Trimming Size",    "Trimming allowance",                  "mm"],
        ["Width Shrinkage",  "Sleeve shrinkage %",                  "%"],
        ["Actual Width",     "Job Width - Trimming (auto-calc)",    "mm"],
        ["UPS",              "Units per sheet (from plan)",         "Number"],
        ["Repeat Length",    "Cylinder repeat (from plan)",         "mm"],
        ["No. of Colors",    "Front + Back colors total",           "Number"],
        ["Print Type",       "Surface Print / Reverse Print",       "Select"],
        ["Quantity",         "Order quantity",                      "Kg / Meter"],
        ["Machine",          "Printing machine",                    "Machine Master"],
    ],
    col_widths=[4.5, 8, 5.5]
)

heading3("Sub-section: Ply Configuration")
body("Tab 1 ke andar hi ply structure fill hoti hai (same as Product Catalog):")
bullet("Multiple plies add kar sakte hain — Film, Ink Layer, Adhesive, Lamination, Coating")
bullet("Har ply mein film sub-group, GSM, thickness, rate set karo")
bullet("Consumable items add karo (ink, adhesive, solvent, hardner) with coverage %")
bullet("Dry weight rows bhi add kar sakte hain for verification")

heading2("3.5  Tab 2 — View Plan (Production)")
body("Tab 1 mein machine select karne ke baad yahan production plan dikhta hai:")
bullet("System automatically possible plans generate karta hai (cylinders × machine constraints)")
bullet("Har plan: cylinder circumference, across UPS, repeat UPS, total UPS, required meters")
bullet("Plan select karo → Tab 1 mein UPS aur dimensions auto-fill ho jaate hain")
bullet("Machine aur process assignment bhi yahan hoti hai:")

add_table(
    ["Column", "Matlab"],
    [
        ["Process",        "Rotogravure process select karo Process Master se"],
        ["Machine",        "Specific machine assign karo"],
        ["Charge Unit",    "SQM / Meter / Color / Job / 1000 Pcs"],
        ["Rate",           "Process ka rate (Process Master se auto-fill)"],
        ["Qty",            "Auto-calculate hota hai chargeUnit ke basis pe"],
        ["Setup Charge",   "One-time setup cost (agar process mein enabled hai)"],
        ["Amount",         "Rate × Qty + Setup Charge"],
    ],
    col_widths=[4, 14]
)

heading2("3.6  Tab 3 — Cost Estimation (Sabse Important)")

heading3("A) Multiple Quantity Costing")
body("Ek saath kai quantities ke liye cost compare kar sakte hain:")
bullet("Extra quantities add karo (e.g., 500kg, 1000kg, 2000kg)")
bullet("Har quantity ke liye alag per-meter rate calculate hoga")
bullet("Table mein side-by-side comparison dikha")

heading3("B) Cost Breakdown Formula (Full Detail)")
code_block([
    "1. MATERIAL COST",
    "   Film Cost  = Σ (GSM ÷ 1000) × Area(m²) × Film Rate per Kg",
    "   Consumable = Σ (Effective GSM ÷ 1000) × Area(m²) × Item Rate",
    "   Note: Effective GSM = GSM × (Coverage% ÷ 100)  [agar coverage < 100%]",
    "",
    "2. PROCESS COST",
    "   = Σ (Rate × Auto-Qty + Setup Charge)",
    "   Auto-Qty logic:",
    "     chargeUnit = 'm²'       →  Qty = Area in m²",
    "     chargeUnit = 'Meter'    →  Qty = Order quantity",
    "     chargeUnit = 'Cylinder' →  Qty = No. of Colors",
    "     chargeUnit = '1000 Pcs' →  Qty = Order Qty ÷ 1000",
    "     chargeUnit = 'Job'      →  Qty = 1",
    "",
    "3. CYLINDER COST",
    "   = No. of Colors × Cost per Color",
    "   (Ya manual override set kar sakte hain)",
    "",
    "4. SETUP COST",
    "   = (Setup Time ÷ 60) × Machine Cost per Hour",
    "   (Ya manual override)",
    "",
    "5. PACKING COST",
    "   Box Cost   = Box Rate + (Plugs × Plug Rate) + (Tape × Tape Rate) + (Stretch Film × Rate)",
    "   Per Kg     = Box Cost ÷ (Coils per Box × Coil Weight)",
    "   Packing Cost = Per Kg × Total Quantity",
    "",
    "6. OVERHEAD",
    "   = Subtotal × Overhead%",
    "   Subtotal = Material + Process + Cylinder + Setup + Labour + Transport + Interest + Packing",
    "",
    "7. PROFIT",
    "   = (Subtotal + Overhead) × Profit%",
    "",
    "8. TOTAL",
    "   = Subtotal + Overhead + Profit",
    "",
    "9. PER METER RATE",
    "   = Total ÷ Effective Quantity",
    "   (Agar quantity < Minimum Order → Minimum Order qty use hoti hai rate ke liye)",
    "",
    "10. MARGIN %",
    "    = (Profit ÷ Total) × 100",
    "",
    "11. BREAK-EVEN QTY",
    "    Fixed Cost  = Cylinder + Setup + Overhead + Labour + Transport + Interest + Packing",
    "    Variable    = (Material + Process) ÷ Quantity  [per meter]",
    "    Contribution = Selling Price - Variable Cost",
    "    Break-Even  = Fixed Cost ÷ Contribution",
])

heading3("C) Additional Cost Fields")
add_table(
    ["Field", "Matlab", "Kahan Jaata Hai"],
    [
        ["Labour Cost",          "Monthly salary based ya manual",           "Fixed cost mein"],
        ["Transportation Cost",  "Delivery/freight charges",                  "Fixed cost mein"],
        ["Interest Cost",        "Loan interest based",                       "Fixed cost mein"],
        ["Loan Amount",          "Loan principal",                            "Interest calc mein"],
        ["Interest Rate %",      "Annual interest rate",                      "Interest calc mein"],
        ["Overhead %",           "Overhead percentage (default 12%)",         "Overhead amount"],
        ["Profit %",             "Profit margin (default 15%)",               "Profit amount"],
        ["Minimum Order Value",  "Minimum order qty floor",                   "Rate calculation mein"],
        ["Selling Price",        "Actual customer ko quote kiya price",        "Break-even calc mein"],
    ],
    col_widths=[5, 7, 6]
)

heading3("D) Cylinder Allocation (Estimation mein bhi)")
body("Estimation page mein bhi cylinder planning hoti hai (similar to Product Catalog):")
bullet("Har color ke liye cylinder assign karo")
bullet("New / Existing / Rechromed / Repeat type specify karo")
bullet("Status track karo: Pending → Available → In Use → Under Chrome → Ordered")

heading3("E) Color Shades (Estimation mein)")
body("Color shade card tab estimation mein bhi hoti hai:")
bullet("Har color ke liye ink item select karo Item Master se")
bullet("GSM specify karo")
bullet("Pantone reference aur LAB values fill karo")
bullet("Shade approval status track karo")

heading2("3.7  Product Catalog Se Load Karna")
body("Estimation mein ek special feature hai — existing Product Catalog se data load karo:")
bullet("'Load from Catalog' button click karo")
bullet("Catalog search karo by product name/number")
bullet("Select karo → saara ply structure, dimensions, processes auto-fill ho jaate hain")
bullet("Estimation number generate hoti hai aur catalog linked rehta hai")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 4 — COST ESTIMATION (Extrusion)
# ══════════════════════════════════════════════════════════════

heading1("4. Cost Estimation Page — Extrusion Film Costing")

body("File: app/cost-estimation/page.tsx")
body("Route: /cost-estimation")

heading2("4.1  Yeh Page Kya Karta Hai?")
body("Yeh page specifically extrusion film banane ki raw material cost calculate karta hai. "
     "Gravure estimation mein film ka rate manually enter karte hain, lekin is page pe "
     "recipe ke based pe exact raw material cost nikalta hai — layer-wise, material-wise.")

heading2("4.2  Key Concepts")

heading3("Recipe")
body("Recipe ek film ka formula hai — kitni layers hain, har layer mein kaunsa "
     "raw material kitne % mein use hota hai:")
code_block([
    "Recipe: '3-Layer LLDPE Film'",
    "  Layer 1 (Outer) — 30% of total micron",
    "    LLDPE  60%  @  ₹120/kg,  density 0.92",
    "    LDPE   40%  @  ₹105/kg,  density 0.92",
    "  Layer 2 (Core)  — 40% of total micron",
    "    LLDPE 100%  @  ₹120/kg",
    "  Layer 3 (Inner) — 30% of total micron",
    "    LLDPE  70%,  Masterbatch 30%",
])

heading3("Micron Distribution")
body("Total film thickness (microns) ko recipe ke layer ratio se divide karta hai:")
code_block([
    "Total Micron = 40µ",
    "Layer Ratio  = 30:40:30",
    "",
    "Layer 1 micron = (30 ÷ 100) × 40 = 12µ",
    "Layer 2 micron = (40 ÷ 100) × 40 = 16µ",
    "Layer 3 micron = (30 ÷ 100) × 40 = 12µ",
])

heading2("4.3  Calculation Engine")
code_block([
    "For each Layer:",
    "  Blend Density     = Σ (Material Density × Material %) / 100",
    "  GSM               = Micron × Blend Density",
    "  Consumption/m²    = GSM ÷ 1000  (kg per m²)",
    "  Blend Rate        = Σ (Material Rate × Material %) / 100",
    "  Cost per m²       = Consumption/m² × Blend Rate",
    "",
    "Total Film Cost     = Σ Cost per m² across all layers × Total Area",
    "Cost per Kg         = Total Film Cost ÷ Order Qty (kg)",
])

heading2("4.4  Editable Fields")
add_table(
    ["Field", "Matlab", "Default"],
    [
        ["Recipe",             "Film ka formula select karo",                 "Recipe Master se"],
        ["Roll Master",        "Standard roll specification",                  "Roll Master se"],
        ["Total Micron",       "Film ki total thickness",                      "Recipe se auto"],
        ["Layer Microns",      "Har layer ka individual micron (editable)",    "Ratio se distribute"],
        ["Material %",         "Har material ka % (editable per layer)",       "Recipe se"],
        ["Material Rate",      "Per kg rate (editable)",                       "Recipe se"],
        ["Order Qty (kg)",     "Kitna order hai",                              "Manual input"],
        ["Customer",           "Customer naam",                                "Select"],
    ],
    col_widths=[4.5, 10, 4.5]
)

heading2("4.5  Output — Kya Milta Hai?")
bullet("Layer-wise GSM breakdown")
bullet("Per m² cost each layer ke liye")
bullet("Material-wise quantity required (kg)")
bullet("Total material cost")
bullet("Per kg film cost (useful for Gravure Estimation mein rate fill karne ke liye)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 5 — CONTEXT & SHARED STATE
# ══════════════════════════════════════════════════════════════

heading1("5. Context Providers — Shared Data Kaise Flow Karta Hai?")

body("Yeh teen modules React Context use karte hain data share karne ke liye. "
     "Context matlab ek global state jo multiple pages mein accessible hota hai.")

add_table(
    ["Context", "File", "Kaunsa Data Provide Karta Hai", "Kaun Use Karta Hai"],
    [
        ["useCategories",    "context/CategoriesContext.tsx",    "Product categories list",           "Catalog, Estimation, Masters"],
        ["useProductCatalog","context/ProductCatalogContext.tsx","Saved catalog entries",             "Catalog page, Estimation page"],
        ["useEnquiries",     "context/EnquiryContext.tsx",       "Gravure enquiries list",            "Estimation page"],
    ],
    col_widths=[4, 5.5, 6, 4.5]
)

heading2("5.1  Context Ka Matlab Kya Hai? (Beginner ke liye)")
code_block([
    "// Context provider app/layout.tsx ya parent component mein hota hai:",
    "<CategoriesProvider>",
    "  <ProductCatalogProvider>",
    "    <YourPage />      ← yahan se context access kar sakte ho",
    "  </ProductCatalogProvider>",
    "</CategoriesProvider>",
    "",
    "// Kisi bhi page mein use karo:",
    'import { useProductCatalog } from "@/context/ProductCatalogContext";',
    "",
    "const { catalog, saveCatalogItem, deleteCatalogItem } = useProductCatalog();",
    "// catalog  = sari catalog entries ka array",
    "// saveCatalogItem(item) = naya item save karo",
    "// deleteCatalogItem(id) = item delete karo",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 6 — SHARED COMPONENTS
# ══════════════════════════════════════════════════════════════

heading1("6. Shared Components — Jo Inn Pages Mein Use Hote Hain")

add_table(
    ["Component", "Import Path", "Kya Karta Hai", "Kaun Use Karta Hai"],
    [
        ["DimensionDiagram",    "components/gravure/DimensionDiagram",   "Visual diagram of product dimensions",        "Catalog, Estimation"],
        ["DimensionInputPanel", "components/gravure/DimensionDiagram",   "Input fields for dimensions",                 "Catalog, Estimation"],
        ["PlanViewer",          "components/gravure/PlanViewer",         "Production plan view",                        "Catalog"],
        ["PlanInput",           "components/gravure/PlanViewer",         "Production plan input form",                  "Catalog"],
        ["DataTable",           "components/tables/DataTable",           "Sortable, filterable table",                  "All pages"],
        ["Button",              "components/ui/Button",                  "Common button with variants",                 "All pages"],
        ["Modal",               "components/ui/Modal",                   "Popup dialog",                               "All pages"],
        ["Input/Select",        "components/ui/Input",                   "Form input elements",                         "All pages"],
        ["statusBadge",         "components/ui/Badge",                   "Colored status badge",                        "All pages"],
    ],
    col_widths=[4, 5, 6, 3]
)

heading2("6.1  DimensionDiagram + CONTENT_TYPE_CONFIG")
body("Yeh component product ki visual diagram dikhata hai. CONTENT_TYPE_CONFIG object "
     "define karta hai kaunse content types ke liye kaunse dimension fields show hone chahiye:")
code_block([
    "// Import karo:",
    'import { DimensionDiagram, DimensionInputPanel, DimValues, CONTENT_TYPE_CONFIG }',
    '  from "@/components/gravure/DimensionDiagram";',
    "",
    "// Use karo:",
    "<DimensionDiagram content={form.content} dimValues={dimValues} />",
    "<DimensionInputPanel",
    "  content={form.content}",
    "  dimValues={dimValues}",
    "  onChange={patchDim}",
    "/>",
])

heading2("6.2  DataTable")
body("DataTable ek powerful reusable table hai:")
code_block([
    'import { DataTable, Column } from "@/components/tables/DataTable";',
    "",
    "// Columns define karo:",
    "const columns: Column<MyType>[] = [",
    '  { key: "name",   header: "Name",   sortable: true },',
    '  { key: "status", header: "Status", render: r => statusBadge(r.status) },',
    '  { key: "action", header: "",       render: r => <button>Edit</button> },',
    "];",
    "",
    "// Use karo:",
    "<DataTable columns={columns} data={myData} />",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 7 — DATA TYPES (TypeScript)
# ══════════════════════════════════════════════════════════════

heading1("7. Key Data Types — TypeScript Types")

heading2("7.1  GravureProductCatalog")
code_block([
    "type GravureProductCatalog = {",
    "  id:              string;",
    "  catalogNo:       string;          // Auto-generated (GC-2425-0001)",
    "  productName:     string;",
    "  customerId:      string;",
    "  customerName:    string;",
    "  categoryId:      string;",
    "  categoryName:    string;",
    "  content:         string;          // Label / Sleeve / Pouch",
    "  jobWidth:        number;          // mm",
    "  jobHeight:       number;          // mm",
    "  noOfColors:      number;",
    "  frontColors:     number;",
    "  backColors:      number;",
    "  machineId:       string;",
    "  machineName:     string;",
    "  standardQty:     number;          // meters",
    "  actualWidth:     number;          // after trimming",
    "  trimmingSize:    number;",
    "  widthShrinkage:  number;          // % for sleeve",
    "  secondaryLayers: SecondaryLayer[];// ply structure",
    "  processes:       GravureEstimationProcess[];",
    "  cylinderCostPerColor: number;",
    "  overheadPct:     number;",
    "  profitPct:       number;",
    "  status:          'Pending' | 'Active' | 'Archived';",
    "  rate:            number;          // per meter rate",
    "}",
])

heading2("7.2  SecondaryLayer (Ply)")
code_block([
    "type SecondaryLayer = {",
    "  layerNo:         number;",
    "  plyType:         string;          // Film / Ink Layer / Adhesive / etc.",
    "  itemSubGroup:    string;          // BOPP / PET / LLDPE etc.",
    "  gsm:             number;",
    "  thickness:       number;          // microns",
    "  filmRate?:       number;          // per kg rate",
    "  consumableItems: PlyConsumableItem[];",
    "}",
    "",
    "type PlyConsumableItem = {",
    "  consumableId: string;",
    "  itemGroup:    string;             // Ink / Adhesive / Solvent",
    "  itemName:     string;",
    "  gsm:          number;",
    "  coveragePct?: number;             // default 100",
    "  rate:         number;",
    "}",
])

heading2("7.3  GravureEstimation")
code_block([
    "type GravureEstimation = {",
    "  id:              string;",
    "  estimationNo:    string;          // GE-2425-0001",
    "  date:            string;",
    "  customerId:      string;",
    "  categoryId:      string;",
    "  jobWidth:        number;",
    "  jobHeight:       number;",
    "  noOfColors:      number;",
    "  quantity:        number;",
    "  unit:            string;          // Kg / Meter",
    "  machineId:       string;",
    "  cylinderCostPerColor: number;     // default 3500",
    "  wastagePct:      number;          // default 1%",
    "  overheadPct:     number;          // default 12%",
    "  profitPct:       number;          // default 15%",
    "  sellingPrice:    number;",
    "  minimumOrderValue: number;",
    "  secondaryLayers: SecondaryLayer[];",
    "  processes:       GravureEstimationProcess[];",
    "  // Calculated fields:",
    "  materialCost:    number;",
    "  processCost:     number;",
    "  cylinderCost:    number;",
    "  totalAmount:     number;",
    "  perMeterRate:    number;",
    "  marginPct:       number;",
    "  breakEvenQty:    number;",
    "  status:          'Draft'|'Approved'|'Sent'|'Accepted'|'Rejected';",
    "}",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 8 — COMMON PATTERNS
# ══════════════════════════════════════════════════════════════

heading1("8. Common Patterns Jo Yeh Pages Use Karte Hain")

heading2("8.1  Form State Update Pattern")
body("In pages mein form update karne ka standard pattern:")
code_block([
    "// Generic field updater:",
    "const rf = <K extends keyof GravureProductCatalog>(",
    "  k: K, v: GravureProductCatalog[K]",
    ") => {",
    "  setReplanForm(p => {",
    "    if (!p) return p;",
    "    const next = { ...p, [k]: v };",
    "    // Dependent fields update:",
    "    if (k === 'frontColors' || k === 'backColors') {",
    "      next.noOfColors = next.frontColors + next.backColors;",
    "    }",
    "    return next;",
    "  });",
    "};",
    "",
    "// Use karo:",
    'rf("jobWidth", 250);',
    'rf("frontColors", 4);',
])

heading2("8.2  useMemo for Expensive Calculations")
body("Cost calculations useMemo mein hain taaki bar bar recalculate na ho:")
code_block([
    "const replanCost = useMemo(() => {",
    "  if (!replanForm) return null;",
    "  // ... saari calculation yahan ...",
    "  return { filmCost, processCost, total, perMeter };",
    "}, [replanForm]);  // ← sirf replanForm badlega toh recalculate hoga",
])

heading2("8.3  unwrap() — Triple-Encoded JSON Handle Karna")
code_block([
    "// Kuch API responses mein JSON string ke andar string hoti hai:",
    'function unwrap(raw: any): any {',
    '  let r = raw;',
    '  while (typeof r === "string") { try { r = JSON.parse(r); } catch { break; } }',
    '  return r;',
    '}',
    "",
    "// Use karo:",
    "const data = unwrap(await res.json());",
])

heading2("8.4  getStructureType() — Content Se Structure Auto-Detect")
code_block([
    "function getStructureType(content: string) {",
    '  const c = content.toLowerCase();',
    '  if (c.includes("ldpe") && c.includes("multi")) return "MultiPackShrink";',
    '  if (c.includes("sleeve"))                       return "Sleeve";',
    '  if (c.includes("pouch") || c.includes("standup")) return "Pouch";',
    '  return "Label";',
    "}",
    "",
    "// Auto-trigger when content changes:",
    'if (k === "content") {',
    '  next.structureType = getStructureType(v as string);',
    '}',
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 9 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════

heading1("9. Quick Reference")

heading2("9.1  Product Catalog — Tab Flow")
code_block([
    "Pending Tab → Click 'Create Catalog'",
    "    ↓",
    "Tab 1: Info    → Product Name, Customer, Category, Dimensions, Machine, Colors",
    "    ↓",
    "Tab 2: Planning → Select Production Plan (UPS, Cylinder, Required Meters)",
    "    ↓",
    "Tab 3: Material → Ply Config, Color Shades, Cylinder Allocation, Attachments",
    "    ↓",
    "Save → Moves to 'Processed' Tab",
])

heading2("9.2  Estimation — Cost Calculation Summary")
add_table(
    ["Cost Component", "Formula", "Override Option"],
    [
        ["Material",     "Σ (GSM÷1000 × Area × Rate)",             "No"],
        ["Process",      "Σ (Rate × Auto-Qty + Setup Charge)",      "No"],
        ["Cylinder",     "Colors × Cost per Color",                 "Yes — cylinderCostOverride"],
        ["Setup",        "Setup Time ÷ 60 × Machine Rate/hr",       "Yes — setupCostOverride"],
        ["Packing",      "Box/plug/tape/stretch formula",            "Yes — packingCostOverride"],
        ["Overhead",     "Subtotal × Overhead%",                    "No (adjust %)"],
        ["Profit",       "(Sub + OH) × Profit%",                    "No (adjust %)"],
        ["Per Meter",    "Total ÷ Qty (or Min Order Qty)",           "Selling Price override"],
        ["Break-Even",   "Fixed Cost ÷ Contribution per Meter",     "No"],
    ],
    col_widths=[4, 8, 6]
)

heading2("9.3  Default Values (Estimation Form)")
add_table(
    ["Field", "Default Value", "Badle Kyun"],
    [
        ["Cylinder Cost/Color",    "₹ 3,500",   "Market rate change hone pe"],
        ["Cylinder Rate/Sq.Inch",  "₹ 2.5",     "Cylinder vendor ka rate"],
        ["Overhead %",             "12%",        "Company policy"],
        ["Profit %",               "15%",        "Company policy"],
        ["Wastage %",              "1%",         "Process pe depend"],
        ["Machine Cost/Hour",      "₹ 1,350",   "Machine depreciation/cost"],
        ["Packing Box Rate",       "₹ 80",      "Packaging material rate"],
    ],
    col_widths=[5, 4, 9]
)

# ══════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════

out = "d:/AJ shrink new code/New updated code/AJ_Shrink_Module_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
