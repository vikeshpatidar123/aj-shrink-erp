from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 73, 125))
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 112, 192))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(68, 68, 68))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_after  = Pt(3)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    return p

def code_block(lines):
    """Light grey box for code"""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(20, 20, 20)
        # grey shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F2F2')
        pPr.append(shd)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

def info_box(lines, bg='EBF3FB', border_color='2E74B5'):
    """Blue info box"""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=10)
        run.font.color.rgb = RGBColor(20, 20, 80)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  bg)
        pPr.append(shd)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '2E74B5')
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        fill = 'FFFFFF' if ri % 2 == 0 else 'EBF3FB'
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill)
            tcPr.append(shd)
    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AJ SHRINK ERP")
set_font(run, size=28, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Frontend Developer Documentation")
set_font(run, size=16, color=(0, 112, 192))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Next.js + Backend API Integration Guide")
set_font(run, size=12, italic=True, color=(80, 80, 80))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0  |  April 2026")
set_font(run, size=10, color=(120, 120, 120))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — NEXT.JS KYA HAI
# ══════════════════════════════════════════════════════════════════════════════

heading1("1. Next.js Kya Hai? (Beginner ke liye)")
body("Next.js ek React framework hai jo frontend applications banane ke liye use hota hai. "
     "Normal React mein hume khud routing setup karni padti hai, lekin Next.js mein folders "
     "se automatically routes bante hain.")

heading2("1.1  React vs Next.js — Fark Kya Hai?")
add_table(
    ["Feature", "Normal React", "Next.js (Is Project Mein)"],
    [
        ["Routing",       "Manually react-router-dom lagao",     "Folders se automatic banta hai"],
        ["Pages",         "Ek hi index.html",                    "Har folder = ek page/route"],
        ["Performance",   "Sab browser pe load hota hai",        "Server-side options bhi hain"],
        ["File structure","Khud banana padta hai",               "app/ folder fixed convention hai"],
    ],
    col_widths=[4, 6, 6]
)

heading2("1.2  App Router Concept")
body("Is project mein Next.js 14+ ka App Router use hota hai. Iska matlab:")
bullet("app/ folder ke andar jo bhi folder banate ho, woh URL route ban jata hai")
bullet("Har route mein page.tsx file hona zaroori hai")
bullet("layout.tsx ek wrapper hota hai jo har page ke upar apply hota hai")

doc.add_paragraph()
code_block([
    "app/",
    "├── layout.tsx              ← Root layout (sidebar, fonts — sab pages pe)",
    "├── page.tsx                ← Home page  →  URL: /",
    "├── login/",
    "│   └── page.tsx            ← URL: /login",
    "├── masters/",
    "│   ├── categories/",
    "│   │   └── page.tsx        ← URL: /masters/categories",
    "│   └── employees/",
    "│       └── page.tsx        ← URL: /masters/employees",
    "├── gravure/",
    "│   ├── orders/",
    "│   │   └── page.tsx        ← URL: /gravure/orders",
    "│   └── product-catalog/",
    "│       └── page.tsx        ← URL: /gravure/product-catalog",
    "└── orders/",
    "    └── page.tsx            ← URL: /orders",
])

doc.add_paragraph()
info_box([
    "  RULE:  Folder naam = URL segment.  page.tsx = woh page ka content.",
    "  Example: app/gravure/orders/page.tsx  →  browser mein /gravure/orders kholo",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

heading1("2. Is Project Ki Poori File Structure")

heading2("2.1  Root Level Files")
add_table(
    ["File / Folder", "Kya Kaam Karta Hai"],
    [
        ["app/",             "Saare pages aur routes yahan hain"],
        ["components/",      "Reusable UI pieces (Button, Table, Modal, etc.)"],
        ["lib/",             "Helper functions — auth, styles, code generation"],
        ["data/",            "Static / dummy data (dummyData.ts)"],
        ["public/",          "Images, icons jo seedhe URL se access hon"],
        ["next.config.ts",   "Next.js configuration file"],
        ["package.json",     "Project dependencies aur scripts"],
        ["tsconfig.json",    "TypeScript configuration"],
    ],
    col_widths=[5, 11]
)

heading2("2.2  app/ Folder — Saare Routes")
add_table(
    ["Route (URL)", "File Path", "Purpose"],
    [
        ["/",                        "app/page.tsx",                          "Home / Dashboard redirect"],
        ["/login",                   "app/login/page.tsx",                    "Login page"],
        ["/dashboard",               "app/dashboard/page.tsx",                "Main dashboard"],
        ["/masters/categories",      "app/masters/categories/page.tsx",       "Category master CRUD"],
        ["/masters/employees",       "app/masters/employees/page.tsx",        "Employee master"],
        ["/masters/hsn",             "app/masters/hsn/page.tsx",              "HSN code master"],
        ["/masters/departments",     "app/masters/departments/page.tsx",      "Department master"],
        ["/masters/item-groups",     "app/masters/item-groups/page.tsx",      "Item group master"],
        ["/gravure/orders",          "app/gravure/orders/page.tsx",           "Gravure orders"],
        ["/gravure/product-catalog", "app/gravure/product-catalog/page.tsx",  "Product catalog"],
        ["/gravure/estimation",      "app/gravure/estimation/page.tsx",       "Gravure estimation"],
        ["/orders",                  "app/orders/page.tsx",                   "Main orders page"],
        ["/production",              "app/production/page.tsx",               "Production tracking"],
        ["/inventory",               "app/inventory/page.tsx",                "Inventory management"],
        ["/dispatch",                "app/dispatch/page.tsx",                 "Dispatch management"],
    ],
    col_widths=[5, 7, 5]
)

heading2("2.3  lib/ Folder — Helper Files")
add_table(
    ["File", "Kya Kaam Karta Hai"],
    [
        ["lib/auth.ts",         "Login functions + authHeaders() — har API call yahi use karta hai"],
        ["lib/styles.ts",       "Common CSS class strings (reuse ke liye)"],
        ["lib/generateCode.ts", "Auto code generation logic (IDs, order numbers, etc.)"],
    ],
    col_widths=[5, 11]
)

heading2("2.4  components/ Folder — Reusable UI Parts")
body("Components woh UI pieces hain jo baar baar alag alag pages pe use hote hain:")
bullet("components/tables/DataTable.tsx  — Table component jo saare pages pe data dikhata hai")
bullet("components/ui/Button.tsx         — Common button with variants")
bullet("components/layout/AppShell.tsx   — Sidebar + main layout wrapper")
bullet("components/ui/Modal.tsx          — Popup/dialog component")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — BACKEND CONNECTION
# ══════════════════════════════════════════════════════════════════════════════

heading1("3. Frontend → Backend Connection Kaise Kaam Karta Hai")

heading2("3.1  Architecture Overview")
body("Is project mein backend ek alag Express.js server hai jo Vercel pe deploy hai. "
     "Frontend (Next.js) seedhe browser se us server ko fetch() calls karta hai.")

doc.add_paragraph()
code_block([
    "Browser (Next.js React)                    Backend (Express.js)",
    "                                           https://api.indusanalytics.co.in",
    "                                           ",
    "  fetch('/api/categorymaster/GetAll')  ──►  Express Route Handler",
    "    headers: authHeaders()                   ↓ Database Query",
    "                                           ◄── JSON Response",
])

doc.add_paragraph()
info_box([
    "  Backend URL:  https://api.indusanalytics.co.in",
    "  Sab API calls yahan jaati hain — koi bhi page banao, yahi BASE_URL use karo",
])

heading2("3.2  Backend Kaise Bana Hai")
body("Backend folder structure:")
code_block([
    "backend/",
    "├── api/",
    "│   ├── index.js       ← Express server entry point",
    "│   └── routes/        ← Alag alag modules ke routes",
    "├── package.json       ← express, cors, dotenv",
    "└── vercel.json        ← Vercel deployment config",
])

heading2("3.3  Authentication Flow — 2 Steps")
body("Har koi directly API use nahi kar sakta. Pehle login karna padta hai. "
     "Login 2 steps mein hota hai:")

heading3("Step 1 — Company Login")
code_block([
    "POST  https://api.indusanalytics.co.in/api/Login/CompanyLogin",
    "",
    "Header:",
    "  Authorization: Basic base64(ApiCompanyUserName:ApiCompanyPassword)",
    "",
    "Response:",
    "  { CompanyName: 'AJ Shrink' }",
    "",
    "Baad mein:  basicAuth  localStorage mein save hota hai",
])

heading3("Step 2 — User Login")
code_block([
    "POST  https://api.indusanalytics.co.in/api/Login/UserLogin",
    "",
    "Header:",
    "  Authorization: Basic base64(ApiCompanyUserName:ApiCompanyPassword)",
    "",
    "Body:",
    "  { UserName: 'admin', Password: 'xxx' }",
    "",
    "Response:",
    "  { UserID, UserName, CompanyID, FYear, DBType, ProductionUnits[] }",
])

heading3("Login ke baad LocalStorage mein save hota hai:")
add_table(
    ["Key", "Value", "Kahan Use Hota Hai"],
    [
        ["basicAuth",        "base64(user:pass)",          "Har API call ke Authorization header mein"],
        ["userID",           "e.g. '5'",                   "API headers mein UserID"],
        ["companyID",        "e.g. '1'",                   "API headers mein CompanyID"],
        ["fYear",            "e.g. '2425'",                "Financial year filtering"],
        ["dbType",           "e.g. 'MSSQL'",               "Database type identify karne"],
        ["productionUnitID", "e.g. '1'",                   "Production unit filtering"],
    ],
    col_widths=[4, 4, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — authHeaders() FUNCTION
# ══════════════════════════════════════════════════════════════════════════════

heading1("4. authHeaders() — Sabse Important Function")

body("lib/auth.ts mein yeh function hai jo har API call ke saath use hota hai:")

code_block([
    "// lib/auth.ts",
    "export function authHeaders(): Record<string, string> {",
    "  return {",
    '    "Content-Type":   "application/json",',
    '    Authorization:    `Basic ${localStorage.getItem("basicAuth") || ""}`,',
    '    CompanyID:        localStorage.getItem("companyID") || "",',
    '    UserID:           localStorage.getItem("userID") || "",',
    '    FYear:            localStorage.getItem("fYear") || "",',
    '    DBType:           localStorage.getItem("dbType") || "MSSQL",',
    '    ProductionUnitID: localStorage.getItem("productionUnitID") || "",',
    "  };",
    "}",
])

doc.add_paragraph()
body("Is function ko kaise use karte hain — example:")

code_block([
    "import { authHeaders } from '@/lib/auth';",
    "",
    "const response = await fetch('https://api.indusanalytics.co.in/api/categorymaster/GetAll', {",
    "  method: 'GET',",
    "  headers: authHeaders()   // ← bas yahi ek line, sab headers automatic",
    "});",
    "",
    "const data = await response.json();",
])

doc.add_paragraph()
info_box([
    "  IMPORTANT: Kabhi bhi hard-code mat karo Authorization headers.",
    "  Hamesha authHeaders() use karo — ek jagah se sab manage hoga.",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — NAYA PAGE KAISE BANATE HAIN
# ══════════════════════════════════════════════════════════════════════════════

heading1("5. Naya Page Kaise Banate Hain — Step by Step")

heading2("5.1  Simple Steps")

body("Example: /masters/suppliers banana hai")

heading3("Step 1 — Folder aur file banao")
code_block([
    "app/masters/suppliers/page.tsx     ← yeh file banao",
])

heading3("Step 2 — Basic template")
code_block([
    '"use client";                           // ← React hooks use karne hain toh ye zaroori',
    'import { useState, useEffect } from "react";',
    'import { authHeaders } from "@/lib/auth";',
    "",
    'const BASE_URL = "https://api.indusanalytics.co.in";',
    "",
    "export default function SuppliersPage() {",
    "  const [suppliers, setSuppliers] = useState([]);",
    "  const [loading, setLoading]     = useState(true);",
    "",
    "  useEffect(() => {",
    "    fetchSuppliers();",
    "  }, []);",
    "",
    "  async function fetchSuppliers() {",
    "    setLoading(true);",
    "    const res  = await fetch(`${BASE_URL}/api/suppliermaster/GetAll`, {",
    "      headers: authHeaders()",
    "    });",
    "    const data = await res.json();",
    "    setSuppliers(data);",
    "    setLoading(false);",
    "  }",
    "",
    "  return (",
    "    <div>",
    '      <h1>Suppliers</h1>',
    "      {loading ? <p>Loading...</p> : (",
    "        <ul>",
    "          {suppliers.map((s: any) => (",
    "            <li key={s.SupplierID}>{s.SupplierName}</li>",
    "          ))}",
    "        </ul>",
    "      )}",
    "    </div>",
    "  );",
    "}",
])

heading2("5.2  Common API Patterns")
add_table(
    ["Operation", "HTTP Method", "URL Pattern", "Code Example"],
    [
        ["Data fetch (list)", "GET",    "/api/xxxmaster/GetAll",      "fetch(url, { headers: authHeaders() })"],
        ["Single item get",  "GET",    "/api/xxxmaster/GetByID/123", "fetch(url, { headers: authHeaders() })"],
        ["Naya add karna",   "POST",   "/api/xxxmaster/Save",        "fetch(url, { method:'POST', headers, body: JSON.stringify(data) })"],
        ["Update karna",     "PUT",    "/api/xxxmaster/Update",      "fetch(url, { method:'PUT', headers, body: JSON.stringify(data) })"],
        ["Delete karna",     "DELETE", "/api/xxxmaster/Delete/123",  "fetch(url, { method:'DELETE', headers })"],
    ],
    col_widths=[3.5, 3, 4.5, 6]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — EXISTING PAGE KA PATTERN
# ══════════════════════════════════════════════════════════════════════════════

heading1("6. Existing Page Ka Pattern — Categories Page Example")

body("app/masters/categories/page.tsx dekho — yeh ek complete CRUD page hai. "
     "Iska pattern samajh lo toh baaki sab pages bana sakte ho:")

add_table(
    ["Part", "Kya Hai", "Code Line"],
    [
        ["BASE_URL",      "Backend server address",                "const BASE_URL = 'https://api.indusanalytics.co.in'"],
        ["BASE",          "Is module ki base URL",                 "const BASE = `${BASE_URL}/api/categorymaster`"],
        ["useState",      "Page ke data/state store karna",       "const [rows, setRows] = useState([])"],
        ["useEffect",     "Page load pe data fetch karna",        "useEffect(() => { fetchData() }, [])"],
        ["authHeaders()", "Har API call mein headers lagana",     "headers: authHeaders()"],
        ["unwrap()",      "Triple-encoded JSON handle karna",     "Custom function — API kabhi kabhi string mein string bhejta hai"],
    ],
    col_widths=[3, 5, 9]
)

heading2("6.1  CRUD Operations Pattern")
code_block([
    "// ── GET (fetch list) ─────────────────────────────────────────────────────",
    "const res  = await fetch(`${BASE}/GetAll`, { headers: authHeaders() });",
    "const data = await res.json();",
    "",
    "// ── POST (create new) ───────────────────────────────────────────────────",
    "const res = await fetch(`${BASE}/Save`, {",
    "  method:  'POST',",
    "  headers: authHeaders(),",
    "  body:    JSON.stringify({ CategoryName: 'New Cat', ... })",
    "});",
    "",
    "// ── PUT (update) ────────────────────────────────────────────────────────",
    "const res = await fetch(`${BASE}/Update`, {",
    "  method:  'PUT',",
    "  headers: authHeaders(),",
    "  body:    JSON.stringify({ CategoryID: '5', CategoryName: 'Updated' })",
    "});",
    "",
    "// ── DELETE ──────────────────────────────────────────────────────────────",
    "const res = await fetch(`${BASE}/Delete/${id}`, {",
    "  method:  'DELETE',",
    "  headers: authHeaders()",
    "});",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — IMPORTANT CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

heading1("7. Important Next.js Concepts Jo Is Project Mein Use Hain")

heading2("7.1  'use client' — Kab Lagate Hain?")
body("Next.js mein by default sab Server Components hote hain. "
     "Agar tum useState, useEffect, onClick jaise browser-side features use karna chahte ho "
     "toh file ke bilkul top pe yeh likhna padta hai:")
code_block([
    '"use client";    // ← File ki pehli line honi chahiye',
    "",
    "// Ab tum ye use kar sakte ho:",
    "import { useState, useEffect } from 'react';",
])

doc.add_paragraph()
add_table(
    ["Situation", "use client chahiye?"],
    [
        ["useState / useEffect use karna",      "Haan — zaroori"],
        ["onClick, onChange events",             "Haan — zaroori"],
        ["Browser APIs (localStorage, window)",  "Haan — zaroori"],
        ["Sirf HTML/JSX return karna",           "Nahi — optional"],
    ],
    col_widths=[8, 8]
)

heading2("7.2  @/ Path Alias")
body("Is project mein @/ ka matlab project root hai. Isse long relative paths se bachte hain:")
code_block([
    "// Bina alias (confusing):",
    "import Button from '../../../components/ui/Button';",
    "",
    "// Alias ke saath (clean):",
    "import Button from '@/components/ui/Button';",
    "import { authHeaders } from '@/lib/auth';",
])

heading2("7.3  layout.tsx — Ek Baar, Har Jagah")
body("app/layout.tsx ek baar define hota hai aur automatically har page ke around wrap hota hai. "
     "Is project mein AppShell component yahan import hai jo sidebar aur navigation provide karta hai. "
     "Tumhe khud sidebar nahi banana — woh automatically aata hai.")

heading2("7.4  TypeScript Types")
body("Is project mein TypeScript use hota hai. Iska matlab har variable ka type pehle define karte hain:")
code_block([
    "// Type define karo",
    "type SupplierRow = {",
    "  SupplierID:   string;",
    "  SupplierName: string;",
    "  City:         string;",
    "};",
    "",
    "// Use karo",
    "const [suppliers, setSuppliers] = useState<SupplierRow[]>([]);",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — COMMON ERRORS
# ══════════════════════════════════════════════════════════════════════════════

heading1("8. Common Errors Aur Unke Solutions")

add_table(
    ["Error", "Reason", "Fix"],
    [
        [
            "localStorage is not defined",
            "Server component mein localStorage use kiya",
            '"use client" file ke top pe lagao'
        ],
        [
            "401 Unauthorized from API",
            "authHeaders() nahi lagaya ya login nahi kiya",
            "authHeaders() import karo aur headers mein pass karo"
        ],
        [
            "CORS error",
            "Backend pe CORS allow nahi",
            "Backend me cors middleware laga hai — check karo origin allow hai ya nahi"
        ],
        [
            "Hydration error",
            "Server aur client ka render alag hai",
            '"use client" lagao ya useEffect mein localStorage access karo'
        ],
        [
            "Module not found: @/components/...",
            "tsconfig.json mein paths nahi set",
            "tsconfig.json mein paths: { \"@/*\": [\"./*\"] } hona chahiye"
        ],
    ],
    col_widths=[5, 6, 6]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

heading1("9. Quick Reference Card")

heading2("9.1  Naya Page Checklist")
bullet("1. app/module-naam/page.tsx file banao")
bullet('2. "use client"; pehli line mein likho')
bullet('3. import { authHeaders } from "@/lib/auth"; add karo')
bullet("4. const BASE_URL = 'https://api.indusanalytics.co.in'; define karo")
bullet("5. useState se state manage karo")
bullet("6. useEffect mein data fetch karo")
bullet("7. Har fetch mein headers: authHeaders() pass karo")

heading2("9.2  Minimal Working Page Template")
code_block([
    '"use client";',
    'import { useState, useEffect } from "react";',
    'import { authHeaders } from "@/lib/auth";',
    "",
    'const BASE_URL = "https://api.indusanalytics.co.in";',
    "",
    "export default function MyPage() {",
    "  const [data, setData]       = useState<any[]>([]);",
    "  const [loading, setLoading] = useState(true);",
    "",
    "  useEffect(() => {",
    "    fetch(`${BASE_URL}/api/mymodule/GetAll`, { headers: authHeaders() })",
    "      .then(r => r.json())",
    "      .then(d => { setData(d); setLoading(false); });",
    "  }, []);",
    "",
    "  if (loading) return <p>Loading...</p>;",
    "",
    "  return (",
    "    <div>",
    "      {data.map((item, i) => <div key={i}>{JSON.stringify(item)}</div>)}",
    "    </div>",
    "  );",
    "}",
])

heading2("9.3  Important Paths")
add_table(
    ["What", "Where"],
    [
        ["Auth function",        "lib/auth.ts  →  authHeaders(), getSession(), loginUser()"],
        ["Backend API URL",      "https://api.indusanalytics.co.in"],
        ["Common UI components", "components/ui/  aur  components/tables/"],
        ["App layout/sidebar",   "components/layout/AppShell.tsx"],
        ["Page routes",          "app/ folder ke andar"],
        ["Static data",          "data/dummyData.ts"],
    ],
    col_widths=[4, 12]
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════

out = "d:/AJ shrink new code/New updated code/AJ_Shrink_ERP_Documentation.docx"
doc.save(out)
print(f"Document saved: {out}")
